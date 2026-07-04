from pathlib import Path
from typing import Callable, Dict, List


SUPPORTED_EXTENSIONS = {".md", ".txt", ".docx", ".xlsx"}


def parse_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def parse_docx_file(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append("[Table %s]" % table_index)
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def parse_xlsx_file(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    parts: List[str] = []

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        parts.append("[Sheet: %s]" % sheet_name)
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                parts.append("%s!R%s: %s" % (sheet_name, row_index, " | ".join(values)))

    workbook.close()
    return "\n".join(parts)


PARSERS: Dict[str, Callable[[Path], str]] = {
    ".md": parse_text_file,
    ".txt": parse_text_file,
    ".docx": parse_docx_file,
    ".xlsx": parse_xlsx_file,
}


def parse_document(path: Path) -> str:
    parser = PARSERS.get(path.suffix.lower())
    if parser is None:
        raise ValueError("Unsupported file type: %s" % path.suffix)
    return parser(path)
